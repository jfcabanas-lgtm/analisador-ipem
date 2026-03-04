# ============================================
# DESPACHO AUDIT - IPEM/RJ
# VERSÃO COM BUSCA REFORÇADA DE DOCUMENTOS
# ============================================

import streamlit as st
import pdfplumber
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import io
from PIL import Image

# CONFIGURAÇÃO DA PÁGINA
st.set_page_config(
    page_title="IPEM - Despacho Inteligente",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# CSS PERSONALIZADO
# ============================================

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Montserrat', sans-serif;
    }
    
    .main-header {
        background: linear-gradient(135deg, #001529 0%, #003366 50%, #0047ab 100%);
        padding: 2.5rem 2rem;
        border-radius: 30px;
        margin-bottom: 2rem;
        box-shadow: 0 20px 40px rgba(0,20,50,0.3);
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    .main-header::before {
        content: "";
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: url('data:image/svg+xml;utf8,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" preserveAspectRatio="none"><polygon points="0,0 100,0 50,100 0,100" fill="rgba(255,255,255,0.03)"/></svg>');
        background-size: 50px 50px;
        opacity: 0.1;
        pointer-events: none;
    }
    
    .header-content {
        display: flex;
        align-items: center;
        justify-content: space-between;
        position: relative;
        z-index: 2;
    }
    
    .header-title {
        flex: 1;
    }
    
    .header-title h1 {
        color: white;
        font-size: 3.5rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: -0.5px;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    
    .header-title h3 {
        color: rgba(255,255,255,0.9);
        font-size: 1.5rem;
        font-weight: 300;
        margin: 0.5rem 0 0 0;
        letter-spacing: 1px;
    }
    
    .header-seal {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        padding: 1.2rem 2rem;
        border-radius: 20px;
        text-align: center;
        border: 2px solid #ffd700;
        box-shadow: 0 10px 20px rgba(0,0,0,0.2);
        min-width: 250px;
    }
    
    .header-seal h2 {
        color: #003366;
        font-size: 1.8rem;
        font-weight: 700;
        margin: 0;
        line-height: 1.2;
    }
    
    .header-seal p {
        color: #666;
        font-size: 1rem;
        margin: 0.5rem 0 0 0;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    .header-seal .sei-link {
        display: inline-block;
        margin-top: 10px;
        padding: 5px 15px;
        background: #ffd700;
        color: #003366;
        text-decoration: none;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    
    .header-seal .sei-link:hover {
        background: #ffffff;
        transform: scale(1.05);
    }
    
    .stats-container {
        display: flex;
        gap: 1.5rem;
        margin: 2rem 0;
    }
    
    .stat-card {
        background: white;
        padding: 1.5rem;
        border-radius: 20px;
        box-shadow: 0 10px 30px rgba(0,51,102,0.1);
        flex: 1;
        text-align: center;
        border: 1px solid #e9ecef;
        transition: transform 0.3s ease;
    }
    
    .stat-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 40px rgba(0,51,102,0.15);
    }
    
    .stat-icon {
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
    }
    
    .stat-label {
        color: #6c757d;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .stat-value {
        color: #003366;
        font-size: 2rem;
        font-weight: 700;
        margin: 0.5rem 0 0 0;
    }
    
    .section-premium {
        background: white;
        padding: 2rem;
        border-radius: 20px;
        margin: 2rem 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.05);
        border-left: 5px solid #003366;
        position: relative;
    }
    
    .section-premium::before {
        content: "⚖️";
        position: absolute;
        top: -15px;
        left: 30px;
        background: #003366;
        color: white;
        font-size: 1.5rem;
        padding: 0.5rem 1rem;
        border-radius: 50px;
        box-shadow: 0 5px 15px rgba(0,51,102,0.3);
    }
    
    .section-title-premium {
        color: #003366;
        font-size: 1.8rem;
        font-weight: 600;
        margin-bottom: 1.5rem;
        padding-left: 3rem;
    }
    
    .upload-premium {
        border: 3px dashed #003366;
        border-radius: 30px;
        padding: 3rem;
        text-align: center;
        background: linear-gradient(135deg, rgba(0,51,102,0.02) 0%, rgba(0,71,171,0.02) 100%);
        transition: all 0.3s ease;
        margin: 2rem 0;
    }
    
    .upload-premium:hover {
        border-color: #0047ab;
        background: linear-gradient(135deg, rgba(0,51,102,0.05) 0%, rgba(0,71,171,0.05) 100%);
        transform: scale(1.02);
    }
    
    .upload-premium span {
        font-size: 4rem;
        display: block;
        margin-bottom: 1rem;
    }
    
    .upload-premium h3 {
        color: #003366;
        font-size: 1.8rem;
        font-weight: 600;
        margin: 0.5rem 0;
    }
    
    .upload-premium p {
        color: #6c757d;
        font-size: 1.1rem;
        max-width: 600px;
        margin: 1rem auto;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #003366 0%, #0047ab 100%);
        color: white;
        font-weight: 600;
        padding: 1rem 2rem;
        border-radius: 15px;
        border: none;
        box-shadow: 0 10px 20px rgba(0,51,102,0.3);
        transition: all 0.3s ease;
        font-size: 1.2rem;
        letter-spacing: 0.5px;
        width: 100%;
        margin: 0.5rem 0;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #0047ab 0%, #003366 100%);
        box-shadow: 0 15px 30px rgba(0,51,102,0.4);
        transform: translateY(-2px);
    }
    
    .footer-premium {
        background: linear-gradient(135deg, #001529 0%, #003366 100%);
        padding: 2rem;
        border-radius: 30px 30px 0 0;
        margin-top: 3rem;
        color: white;
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    
    .footer-premium::before {
        content: "";
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, #ffd700, #ffffff, #ffd700);
    }
    
    .footer-premium p {
        margin: 0.5rem 0;
        font-size: 1rem;
        opacity: 0.9;
    }
    
    .footer-premium strong {
        color: #ffd700;
        font-weight: 600;
    }
    
    .footer-premium a {
        color: #ffd700;
        text-decoration: none;
        font-weight: 600;
        padding: 5px 20px;
        border: 1px solid #ffd700;
        border-radius: 20px;
        transition: all 0.3s ease;
        display: inline-block;
        margin: 1rem 0;
    }
    
    .footer-premium a:hover {
        background: #ffd700;
        color: #003366;
    }
    
    .timeline {
        display: flex;
        justify-content: space-between;
        margin: 2rem 0;
        position: relative;
    }
    
    .timeline::before {
        content: "";
        position: absolute;
        top: 30px;
        left: 50px;
        right: 50px;
        height: 3px;
        background: linear-gradient(90deg, #003366, #0047ab, #003366);
        z-index: 1;
    }
    
    .timeline-step {
        text-align: center;
        position: relative;
        z-index: 2;
        flex: 1;
    }
    
    .timeline-icon {
        background: white;
        width: 60px;
        height: 60px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        margin: 0 auto 1rem;
        border: 3px solid #003366;
        box-shadow: 0 5px 15px rgba(0,51,102,0.2);
        font-size: 1.8rem;
    }
    
    .timeline-label {
        background: #f8f9fa;
        padding: 0.5rem 1rem;
        border-radius: 30px;
        font-weight: 600;
        color: #003366;
        font-size: 0.9rem;
        display: inline-block;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    .success-premium {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        color: #155724;
        padding: 1.5rem;
        border-radius: 15px;
        border-left: 5px solid #28a745;
        margin: 1rem 0;
        font-size: 1.1rem;
        font-weight: 500;
        box-shadow: 0 5px 15px rgba(40,167,69,0.2);
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# HEADER PRINCIPAL COM LINK DO SEI
# ============================================

st.markdown("""
<div class="main-header">
    <div class="header-content">
        <div class="header-title">
            <h1>IPEM/RJ</h1>
            <h3>INSTITUTO DE PESOS E MEDIDAS DO ESTADO DO RIO DE JANEIRO</h3>
        </div>
        <div class="header-seal">
            <h2>AUDITORIA<br>INTERNA</h2>
            <p>ANÁLISE DE PROCESSO DE LICITAÇÃO E DISPENSA</p>
            <a href="https://sei.rj.gov.br/sei/" target="_blank" class="sei-link">
                🔐 ACESSAR SEI
            </a>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# ============================================
# MENSAGEM DE BOAS-VINDAS
# ============================================

st.markdown("""
<div class="section-premium">
    <div class="section-title-premium">Bem-vindo à Análise de Processos de Licitação e Dispensa</div>
    <p style="font-size: 1.2rem; color: #495057; line-height: 1.6;">
        Este sistema foi desenvolvido para automatizar a geração de despachos de auditoria,
        garantindo padronização, agilidade e conformidade com a Lei nº 14.133/2021.
        Através da análise inteligente de documentos, você poderá gerar despachos
        completos em poucos minutos.
    </p>
</div>
""", unsafe_allow_html=True)

# ============================================
# TIMELINE DO PROCESSO
# ============================================

st.markdown("""
<div class="timeline">
    <div class="timeline-step">
        <div class="timeline-icon">📂</div>
        <span class="timeline-label">Upload do PDF</span>
    </div>
    <div class="timeline-step">
        <div class="timeline-icon">🔍</div>
        <span class="timeline-label">Análise Automática</span>
    </div>
    <div class="timeline-step">
        <div class="timeline-icon">✏️</div>
        <span class="timeline-label">Confirmação de Dados</span>
    </div>
    <div class="timeline-step">
        <div class="timeline-icon">📥</div>
        <span class="timeline-label">Download do Despacho</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ============================================
# INICIALIZAÇÃO DO SESSION STATE
# ============================================

if 'dados_extraidos' not in st.session_state:
    st.session_state.dados_extraidos = None
if 'texto_extraido' not in st.session_state:
    st.session_state.texto_extraido = None
if 'seis_encontrados' not in st.session_state:
    st.session_state.seis_encontrados = []
if 'doc_bytes' not in st.session_state:
    st.session_state.doc_bytes = None
if 'nome_arquivo' not in st.session_state:
    st.session_state.nome_arquivo = None
if 'processos_analisados' not in st.session_state:
    st.session_state.processos_analisados = 0

# ============================================
# ESTATÍSTICAS
# ============================================

if not st.session_state.dados_extraidos and not st.session_state.doc_bytes:
    
    st.markdown(f"""
    <div class="stats-container">
        <div class="stat-card">
            <div class="stat-icon">⚡</div>
            <div class="stat-label">Processos Analisados</div>
            <div class="stat-value">{st.session_state.processos_analisados}</div>
        </div>
        <div class="stat-card">
            <div class="stat-icon">📊</div>
            <div class="stat-label">Despachos Gerados</div>
            <div class="stat-value">{st.session_state.processos_analisados}</div>
        </div>
        <div class="stat-card">
            <div class="stat-icon">⏱️</div>
            <div class="stat-label">Tempo Médio</div>
            <div class="stat-value">2 min</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ============================================
# PASSO 1: UPLOAD DO PDF
# ============================================

st.markdown("""
<div class="upload-premium">
    <span>📎</span>
    <h3>Upload do Processo</h3>
    <p>Arraste o arquivo PDF do processo SEI ou clique para selecionar</p>
</div>
""", unsafe_allow_html=True)

arquivo = st.file_uploader("", type=['pdf'], label_visibility="collapsed")

# ============================================
# FUNÇÕES REFORÇADAS PARA EXTRAIR DADOS
# ============================================

def extrair_campo(padroes, texto, default=""):
    for padrao in padroes:
        match = re.search(padrao, texto, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return default

def encontrar_tr(texto):
    """Busca por TR em qualquer formato no texto"""
    
    # Lista de PADRÕES REFORÇADOS para encontrar o TR
    padroes_tr = [
        r'TR[:\s]*n[º°]?\s*(\d+[/-]\d+)',
        r'TR[:\s]*n[º°]?\s*(\d+)',
        r'Termo de Referência[:\s]*n[º°]?\s*(\d+[/-]\d+)',
        r'Termo de Referência[:\s]*n[º°]?\s*(\d+)',
        r'TR[-]?(\d+[/-]\d+)',
        r'TR[-]?(\d+)',
        r'TR[:\s]*(\d+)',
        r'TR[:\s]*(\d+/\d+)',
        r'T\.R\.\s*n[º°]?\s*(\d+/\d+)',
        r'T\.R\.\s*n[º°]?\s*(\d+)',
        r'(\d{1,3}[/-]\d{4})'  # Qualquer número no formato XX/XXXX
    ]
    
    for padrao in padroes_tr:
        matches = re.findall(padrao, texto, re.IGNORECASE)
        if matches:
            return matches[0]
    return ""

def encontrar_etp(texto):
    """Busca por ETP em qualquer formato no texto"""
    padroes_etp = [
        r'ETP[:\s]*n[º°]?\s*(\d+[/-]\d+)',
        r'ETP[:\s]*n[º°]?\s*(\d+)',
        r'Estudo Técnico Preliminar[:\s]*n[º°]?\s*(\d+[/-]\d+)',
        r'Estudo Técnico Preliminar[:\s]*n[º°]?\s*(\d+)',
        r'ETP[-]?(\d+[/-]\d+)',
        r'ETP[-]?(\d+)',
        r'E\.T\.P\.\s*n[º°]?\s*(\d+/\d+)',
        r'E\.T\.P\.\s*n[º°]?\s*(\d+)',
    ]
    
    for padrao in padroes_etp:
        matches = re.findall(padrao, texto, re.IGNORECASE)
        if matches:
            return matches[0]
    return ""

def encontrar_seis(texto):
    """Busca TODOS os números SEI no texto"""
    padroes_sei = [
        r'SEI[:\s]*n[º°]?\s*(\d+)',
        r'SEI[:\s]*(\d+)',
        r'Documento[:\s]*SEI[:\s]*n[º°]?\s*(\d+)',
        r'Despacho[:\s]*SEI[:\s]*n[º°]?\s*(\d+)',
        r'SEI[-]?(\d+)',
    ]
    
    todos_seis = []
    for padrao in padroes_sei:
        matches = re.findall(padrao, texto, re.IGNORECASE)
        todos_seis.extend(matches)
    
    # Remove duplicatas mantendo a ordem
    seen = set()
    seis_unicos = []
    for sei in todos_seis:
        if sei not in seen:
            seen.add(sei)
            seis_unicos.append(sei)
    
    return seis_unicos

def encontrar_risco(texto):
    """Busca por Matriz de Riscos"""
    padroes_risco = [
        r'Matriz de Riscos[:\s]*n[º°]?\s*(\d+/\d+)',
        r'Matriz de Riscos[:\s]*n[º°]?\s*(\d+)',
        r'Gestão de Risco[:\s]*n[º°]?\s*(\d+/\d+)',
        r'Risco[:\s]*n[º°]?\s*(\d+/\d+)',
    ]
    
    for padrao in padroes_risco:
        matches = re.findall(padrao, texto, re.IGNORECASE)
        if matches:
            return matches[0]
    return ""

# ============================================
# PROCESSAR O PDF
# ============================================

if arquivo and st.session_state.dados_extraidos is None:
    
    with st.spinner("🔍 Analisando PDF e extraindo dados..."):
        
        texto = ""
        with pdfplumber.open(io.BytesIO(arquivo.read())) as pdf:
            for pagina in pdf.pages:
                if pagina.extract_text():
                    texto += pagina.extract_text() + "\n"
        
        st.session_state.texto_extraido = texto
        
        # MOSTRAR PARTE DO TEXTO PARA DIAGNÓSTICO
        with st.expander("📄 Ver texto extraído do PDF (primeiras 1000 caracteres)"):
            st.text(texto[:1000])
        
        # Extrair dados usando as funções reforçadas
        dados_extraidos = {
            'processo_sei': extrair_campo([
                r'Processo[:\s]*n[º°]?\s*([\d\-/]+)',
                r'SEI[:\s]*n[º°]?\s*([\d\-/]+)',
                r'(\d{6,}/\d{6,}/\d{4})'
            ], texto, ""),
            
            'objeto': extrair_campo([
                r'objeto[:\s]*([^.]+)',
                r'aquisição[:\s]*([^.]+)',
                r'contratação[:\s]*([^.]+)'
            ], texto, ""),
            
            'valor': extrair_campo([
                r'R\$\s*([\d.,]+)',
                r'valor[:\s]*R\$\s*([\d.,]+)',
                r'total[:\s]*R\$\s*([\d.,]+)'
            ], texto, ""),
            
            'etp_numero': encontrar_etp(texto),
            'tr_numero': encontrar_tr(texto),
            'risco_numero': encontrar_risco(texto),
            
            'req_siga': extrair_campo([
                r'Requisição[:\s]*n[º°]?\s*(\d+/\d+)',
                r'Requisição SIGA[:\s]*n[º°]?\s*(\d+/\d+)',
                r'SIGA[:\s]*n[º°]?\s*(\d+/\d+)',
                r'Requisição[:\s]*n[º°]?\s*(\d+)'
            ], texto, ""),
            
            'parecer_numero': extrair_campo([
                r'Despacho SEI[:\s]*n[º°]?\s*(\d+)',
                r'Parecer[:\s]*n[º°]?\s*(\d+)',
                r'Parecer Jurídico[:\s]*n[º°]?\s*(\d+)'
            ], texto, ""),
            
            'data_autorizacao': extrair_campo([
                r'autorizado[:\s]*em[:\s]*(\d{1,2}[/]\d{1,2}[/]\d{4})',
                r'(\d{1,2}[/]\d{1,2}[/]\d{4})'
            ], texto, "")
        }
        
        # Encontrar todos os SEIs
        seis_encontrados = encontrar_seis(texto)
        
        # MOSTRAR DIAGNÓSTICO DO QUE FOI ENCONTRADO
        st.info("🔍 **DIAGNÓSTICO DA EXTRAÇÃO:**")
        
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.write(f"**📋 Processo:** {dados_extraidos['processo_sei'] or '❌'}")
            st.write(f"**📝 Objeto:** {dados_extraidos['objeto'][:50] or '❌'}...")
            st.write(f"**💰 Valor:** {dados_extraidos['valor'] or '❌'}")
            
        with col_d2:
            st.write(f"**📄 ETP:** {dados_extraidos['etp_numero'] or '❌ NÃO ENCONTRADO'}")
            st.write(f"**📄 TR:** {dados_extraidos['tr_numero'] or '❌ NÃO ENCONTRADO'}")
            st.write(f"**📄 Parecer:** {dados_extraidos['parecer_numero'] or '❌'}")
        
        st.write(f"**🔢 SEIs encontrados:** {len(seis_encontrados)} números")
        if seis_encontrados:
            st.write(f"   • {', '.join(seis_encontrados[:5])}")
        
        st.session_state.dados_extraidos = dados_extraidos
        st.session_state.seis_encontrados = seis_encontrados
        st.session_state.processos_analisados += 1
        
        st.rerun()

# ============================================
# FUNÇÃO PARA GERAR DESPACHO CONFORME MODELO
# ============================================

def gerar_despacho_modelo(processo_sei, objeto, data_autorizacao, 
                         valor_input, seis, 
                         etp_numero, sei_etp,
                         tr_numero, sei_tr,
                         risco_numero, sei_risco,
                         req_siga, parecer_numero,
                         sei_impacto, sei_disponibilidade, sei_ordenador,
                         fundamentacao, observacoes):
    
    doc = Document()
    
    # CONFIGURAÇÃO DE ESTILOS
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # ========================================
    # I. INTRODUÇÃO
    # ========================================
    
    p = doc.add_paragraph()
    run = p.add_run("I. Introdução")
    run.bold = True
    
    doc.add_paragraph(
        f"Atendendo à solicitação de análise do processo SEI nº {processo_sei}, "
        f"pela Diretoria de Administração e Finanças – DIRAF, referente à {objeto} "
        f"para o Instituto de Pesos e Medidas do Estado do Rio de Janeiro (IPEM/RJ), "
        f"procedemos à verificação dos documentos apresentados, com o objetivo de "
        f"subsidiar a continuidade do processo, sem adentrar no mérito técnico da contratação."
    )
    
    doc.add_paragraph()
    
    # ========================================
    # II. ANÁLISE PROCESSUAL
    # ========================================
    
    p = doc.add_paragraph()
    run = p.add_run("II. Análise Processual")
    run.bold = True
    
    doc.add_paragraph(
        "Foram examinados os seguintes documentos e etapas processuais:"
    )
    
    doc.add_paragraph()
    
    # 1. Solicitação Inicial e Autorização
    p = doc.add_paragraph()
    p.add_run("1. Solicitação Inicial e Autorização: ").bold = True
    
    texto1 = f"O processo foi iniciado pela Superintendência de Pré-Medidos "
    if len(seis) > 0:
        texto1 += f"SEI {seis[0]} "
    texto1 += f"e devidamente autorizado pela Presidência do IPEM/RJ em "
    if data_autorizacao:
        texto1 += f"{data_autorizacao.strftime('%d/%m/%Y')} "
    if len(seis) > 1:
        texto1 += f"SEI {seis[1]}"
    texto1 += "."
    
    p.add_run(texto1)
    
    # 2. ETP e Gestão de Riscos
    p = doc.add_paragraph()
    p.add_run("2. Estudo Técnico Preliminar-ETP e Gestão de Riscos: ").bold = True
    
    texto2 = ""
    if etp_numero:
        texto2 += f"O ETP nº {etp_numero} "
        if sei_etp:
            texto2 += f"SEI {sei_etp} "
    else:
        texto2 += "O ETP não foi identificado na análise automática. "
    
    if risco_numero:
        texto2 += f"e a Matriz de Riscos nº {risco_numero} "
        if sei_risco:
            texto2 += f"SEI {sei_risco} "
    else:
        texto2 += "e a Matriz de Riscos não foi identificada. "
    
    texto2 += "Estes documentos detalham a necessidade, viabilidade técnica e ações de mitigação para a contratação."
    p.add_run(texto2)
    
    # 3. Termo de Referência (VERSÃO MELHORADA)
    p = doc.add_paragraph()
    p.add_run("3. Termo de Referência-TR: ").bold = True
    
    texto3 = ""
    if tr_numero:
        texto3 += f"O TR nº {tr_numero}"
        if sei_tr:
            texto3 += f", SEI {sei_tr}"
        texto3 += ", consolida as especificações técnicas e condições contratuais, servindo de balizador para a fase externa."
    else:
        texto3 += "O TR não foi localizado na análise automática. "
        texto3 += "Recomenda-se verificar a existência deste documento nos autos, "
        texto3 += "pois é essencial para a instrução processual."
        texto3 += "\n  → Caso o documento exista, informe o número no campo acima."
    
    p.add_run(texto3)
    
    # 4. Pesquisa de Mercado e Requisição SIGA
    p = doc.add_paragraph()
    p.add_run("4. Pesquisa de Mercado e Requisição SIGA: ").bold = True
    
    texto4 = "Foi realizada pesquisa de mercado formal, com a devida inclusão da Requisição de Material "
    if req_siga:
        texto4 += f"nº {req_siga} "
    else:
        texto4 += "não identificada "
    
    texto4 += "no Sistema Integrado de Gestão de Aquisições (SIGA)"
    
    if valor_input:
        texto4 += f", totalizando o valor estimado de {valor_input}"
    
    texto4 += "."
    p.add_run(texto4)
    
    # 5. Conformidade Orçamentária
    p = doc.add_paragraph()
    p.add_run("5. Conformidade Orçamentária: ").bold = True
    
    texto5 = "O processo conta com as declarações de impacto financeiro "
    if sei_impacto:
        texto5 += f"SEI {sei_impacto}, "
    else:
        texto5 += "não localizada, "
    
    texto5 += "disponibilidade orçamentária "
    if sei_disponibilidade:
        texto5 += f"SEI {sei_disponibilidade} "
    else:
        texto5 += "não localizada "
    
    texto5 += "e a declaração do ordenador de despesa "
    if sei_ordenador:
        texto5 += f"SEI {sei_ordenador}"
    else:
        texto5 += "não localizada"
    
    texto5 += ", atestando a compatibilidade com o orçamento e o Plano Plurianual (PPA/RJ) para 2026."
    p.add_run(texto5)
    
    # 6. Parecer Jurídico
    p = doc.add_paragraph()
    p.add_run("6. Parecer Jurídico: ").bold = True
    
    texto6 = "A Diretoria Jurídica manifestou-se por meio do Despacho SEI "
    if parecer_numero:
        texto6 += f"{parecer_numero}"
    else:
        texto6 += "não identificado"
    
    texto6 += f", informando a dispensa de análise jurídica formal em razão do valor da contratação, fundamentada em {fundamentacao}."
    
    p.add_run(texto6)
    
    doc.add_paragraph()
    
    # ========================================
    # III. OBSERVAÇÕES
    # ========================================
    
    p = doc.add_paragraph()
    run = p.add_run("III. Observações")
    run.bold = True
    
    if observacoes:
        doc.add_paragraph(observacoes)
    else:
        doc.add_paragraph(
            "Verifica-se que o processo encontra-se devidamente instruído, tendo percorrido as etapas "
            "formais exigidas pela legislação vigente. As especificações técnicas e condições de "
            "fornecimento estão consolidadas no Termo de Referência, documento que orientará a fase de "
            "seleção do fornecedor."
        )
    
    doc.add_paragraph()
    
    # ========================================
    # IV. DESPACHO
    # ========================================
    
    p = doc.add_paragraph()
    run = p.add_run("IV. Despacho")
    run.bold = True
    
    # Verificar se todos os documentos essenciais estão presentes
    docs_essenciais = etp_numero and tr_numero and valor_input and parecer_numero
    
    if docs_essenciais:
        texto_despacho = (
            "Dessa forma, e considerando que os atos administrativos até o presente momento se mostram "
            "formalmente adequados e em conformidade com a Lei nº 14.133/2021 e demais normas aplicáveis, "
            f"indicamos à continuidade do processo de {objeto}."
        )
    else:
        texto_despacho = (
            "Dessa forma, considerando a necessidade de complementação documental, "
            "recomendamos a regularização dos itens pendentes antes do prosseguimento do feito, "
            "conforme observações acima."
        )
    
    doc.add_paragraph(texto_despacho)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========================================
    # ASSINATURA
    # ========================================
    
    doc.add_paragraph("At.te.,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("___________________________________")
    doc.add_paragraph("Auditor Interno")
    doc.add_paragraph("IPEM/RJ")
    
    return doc

# ============================================
# SE HOUVER DADOS EXTRAÍDOS, MOSTRA O FORMULÁRIO
# ============================================

if st.session_state.dados_extraidos:
    
    dados = st.session_state.dados_extraidos
    seis = st.session_state.seis_encontrados
    
    st.markdown("""
    <div class="section-premium">
        <div class="section-title-premium">📊 Dados Encontrados no PDF</div>
    </div>
    """, unsafe_allow_html=True)
    
    col_res1, col_res2, col_res3 = st.columns(3)
    
    with col_res1:
        st.info("📋 **Processo**")
        st.write(f"**Nº:** {dados['processo_sei'] or 'Não encontrado'}")
        st.write(f"**Data:** {dados['data_autorizacao'] or 'Não encontrada'}")
    
    with col_res2:
        st.info("💰 **Valor**")
        if dados['valor']:
            try:
                valor_float = float(dados['valor'].replace('.', '').replace(',', '.'))
                st.write(f"**R$ {valor_float:,.2f}**")
            except:
                st.write(f"**R$ {dados['valor']}**")
        else:
            st.write("**Não identificado**")
    
    with col_res3:
        st.info("📑 **Documentos**")
        st.write(f"ETP: {dados['etp_numero'] or '❌'}")
        st.write(f"TR: {dados['tr_numero'] or '❌'}")
        st.write(f"Parecer: {dados['parecer_numero'] or '❌'}")
    
    if seis:
        with st.expander(f"📎 {len(seis)} números SEI encontrados"):
            for i, sei in enumerate(seis[:10]):
                st.write(f"• SEI {sei}")
    
    # ========================================
    # FORMULÁRIO
    # ========================================
    
    st.markdown("""
    <div class="section-premium">
        <div class="section-title-premium">✏️ Confirmação de Dados</div>
    </div>
    """, unsafe_allow_html=True)
    
    with st.form("form_confirmacao"):
        
        col1, col2 = st.columns(2)
        
        with col1:
            processo_sei = st.text_input("Nº do Processo SEI *", value=dados['processo_sei'])
            objeto = st.text_input("Objeto da Contratação *", value=dados['objeto'])
            
            try:
                data_valor = datetime.strptime(dados['data_autorizacao'], '%d/%m/%Y') if dados['data_autorizacao'] and '/' in dados['data_autorizacao'] else None
            except:
                data_valor = None
            
            data_autorizacao = st.date_input("Data da Autorização", value=data_valor, format="DD/MM/YYYY")
        
        with col2:
            valor_input = st.text_input("Valor (R$)", value=dados['valor'])
            sei_inicio = st.text_input("SEI da Solicitação Inicial", value=seis[0] if len(seis) > 0 else "")
            sei_autorizacao = st.text_input("SEI da Autorização", value=seis[1] if len(seis) > 1 else "")
        
        col3, col4 = st.columns(2)
        
        with col3:
            etp_numero = st.text_input("Nº do ETP", value=dados['etp_numero'])
            sei_etp = st.text_input("SEI do ETP", value=seis[2] if len(seis) > 2 else "")
            tr_numero = st.text_input("Nº do TR *", value=dados['tr_numero'], 
                                     help="Campo obrigatório - informe o número do TR")
            sei_tr = st.text_input("SEI do TR", value=seis[3] if len(seis) > 3 else "")
        
        with col4:
            risco_numero = st.text_input("Nº da Matriz de Riscos", value=dados['risco_numero'])
            sei_risco = st.text_input("SEI da Matriz de Riscos", value=seis[4] if len(seis) > 4 else "")
            req_siga = st.text_input("Nº da Requisição SIGA", value=dados['req_siga'])
            parecer_numero = st.text_input("Nº do Parecer Jurídico", value=dados['parecer_numero'])
        
        col5, col6 = st.columns(2)
        
        with col5:
            sei_impacto = st.text_input("SEI da Declaração de Impacto", value=seis[5] if len(seis) > 5 else "")
            sei_disponibilidade = st.text_input("SEI da Disponibilidade Orçamentária", value=seis[6] if len(seis) > 6 else "")
            sei_ordenador = st.text_input("SEI da Declaração do Ordenador", value=seis[7] if len(seis) > 7 else "")
        
        with col6:
            fundamentacao = st.text_area("Fundamentação Legal", 
                value="art. 1º da Resolução PGE nº 5.059/2024 e no art. 95, I, da Lei nº 14.133/2021",
                height=100)
        
        observacoes = st.text_area("Observações", height=100)
        
        submitted = st.form_submit_button("✅ CONFIRMAR DADOS E GERAR DESPACHO", use_container_width=True)
    
    # ========================================
    # GERAR DESPACHO
    # ========================================
    
    if submitted:
        
        if not processo_sei or not objeto:
            st.error("❌ Processo e Objeto são obrigatórios!")
        else:
            
            with st.spinner("Gerando despacho conforme modelo..."):
                
                # Usar a função com o modelo
                doc = gerar_despacho_modelo(
                    processo_sei, objeto, data_autorizacao, 
                    valor_input, seis, 
                    etp_numero, sei_etp,
                    tr_numero, sei_tr,
                    risco_numero, sei_risco,
                    req_siga, parecer_numero,
                    sei_impacto, sei_disponibilidade, sei_ordenador,
                    fundamentacao, observacoes
                )
                
                # Salvar em memória
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                st.session_state.doc_bytes = doc_bytes.getvalue()
                st.session_state.nome_arquivo = f"DESPACHO_{processo_sei.replace('/', '_')}.docx"
                
                st.rerun()

# ============================================
# DOWNLOAD DO DESPACHO
# ============================================

if st.session_state.doc_bytes:
    
    st.markdown("""
    <div class="section-premium">
        <div class="section-title-premium">✅ Despacho Gerado com Sucesso!</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.balloons()
    
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    
    with col_btn2:
        st.download_button(
            label="📥 BAIXAR DESPACHO",
            data=st.session_state.doc_bytes,
            file_name=st.session_state.nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        if st.button("🔄 NOVO PROCESSO", use_container_width=True):
            for key in ['dados_extraidos', 'texto_extraido', 'seis_encontrados', 'doc_bytes', 'nome_arquivo']:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

# ============================================
# RODAPÉ INSTITUCIONAL
# ============================================

st.markdown("""
<div class="footer-premium">
    <p><strong>IPEM/RJ - INSTITUTO DE PESOS E MEDIDAS DO ESTADO DO RIO DE JANEIRO</strong></p>
    <p>AUDITORIA INTERNA - ANÁLISE DE PROCESSO DE LICITAÇÃO E DISPENSA</p>
    <a href="https://sei.rj.gov.br/sei/" target="_blank">🔐 ACESSAR SEI</a>
    <p style="font-size: 0.9rem; opacity: 0.8;">Lei nº 14.133/2021 • Versão 9.0 - Busca Reforçada</p>
    <p style="font-size: 0.8rem; opacity: 0.6;">© 2026 - Todos os direitos reservados</p>
</div>
""", unsafe_allow_html=True)
